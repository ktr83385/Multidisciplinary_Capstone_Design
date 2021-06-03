import pandas as pd
import matplotlib.pyplot as plt

from influxdb_client import InfluxDBClient, Point
from influxdb_client.client.write_api import SYNCHRONOUS



bucket = "test"

client = InfluxDBClient(url="http://52.79.255.212:8086/", user ="admin", password="1q2w3e4r!", token="dLCN0Dt_TXxk4ASzM6rJMaKa9T2Q6MrW65ayAYht-n69TdRVxuXXSQ_yo-pmfOhIElL7FLbveib-0gtBpkuiIA==", org="4859345f660a00f5")

write_api = client.write_api(write_options=SYNCHRONOUS)
query_api = client.query_api()

"""
Query: using Pandas DataFrame
"""

data_frame_1 = query_api.query_data_frame('from(bucket:"test") '
                                                '|> range(start: 2021-05-18, stop: 2021-05-19)'
                                                '|> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")'
                                                '|> keep(columns: ["_time", "CO2_ppm_1f", "CO_ppm_1f", "PM1.0_1f", "PM10.0_1f", "PM2.5_1f"])'

                                       )
data_frame_2 = query_api.query_data_frame('from(bucket:"test") '
                                                '|> range(start: 2021-05-18, stop: 2021-05-19)'
                                                '|> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")'
                                                '|> keep(columns: ["_time", "CO2_ppm_3f", "CO_ppm_3f", "PM1.0_3f", "PM10.0_3f", "PM2.5_3f" ])'
                                       )
df_1f = data_frame_1.dropna(axis=0)
df_3f = data_frame_2.dropna(axis=0)

#data_frame_3.to_csv('sample_1.csv')
#data_frame_4.to_csv('sample_2.csv')

"""
Close client
"""
client.close()

import numpy as np
df_1f.index = np.arange(len(df_1f))
df_3f.index = np.arange(len(df_3f))

fig, axes = plt.subplots(5,1, figsize=(40,40))
plt.rc('ytick', labelsize=50)

axes[0].set_title("C02 24h curve", fontsize = 40)
axes[0].set_xlabel("Time")
axes[0].set_ylabel("CO2 ppm", fontsize = 30)

axes[0].plot(df_1f['_time'][::600], df_1f['CO2_ppm_1f'][::600], color = 'r', label ='1f')   # 10분 간격으로 수치값 변화
axes[0].plot(df_3f['_time'][::600], df_3f['CO2_ppm_3f'][::600], color = 'b', label = '3f')
axes[0].legend(loc='best')

axes[1].set_title("C0 24h curve", fontsize = 40)
axes[1].set_xlabel("Time")
axes[1].set_ylabel("CO ppm", fontsize = 30)

axes[1].plot(df_1f['_time'][::600], df_1f['CO_ppm_1f'][::600], color = 'r', label ='1f')   # 10분 간격으로 수치값 변화
axes[1].plot(df_3f['_time'][::600], df_3f['CO_ppm_3f'][::600], color = 'b', label = '3f')
axes[1].legend(loc='best')

axes[2].set_title("PM1.0 24h curve", fontsize = 40)
axes[2].set_xlabel("Time")
axes[2].set_ylabel("PM1.0", fontsize = 30)

axes[2].plot(df_1f['_time'][::600], df_1f['PM1.0_1f'][::600], color = 'r', label ='1f')   # 10분 간격으로 수치값 변화
axes[2].plot(df_3f['_time'][::600], df_3f['PM1.0_3f'][::600], color = 'b', label = '3f')
axes[2].legend(loc='best', fontsize = 30)

axes[3].set_title("PM2.5 24h curve", fontsize = 40)
axes[3].set_xlabel("Time")
axes[3].set_ylabel("PM2.5", fontsize = 30)

axes[3].plot(df_1f['_time'][::600], df_1f['PM2.5_1f'][::600], color = 'r', label ='1f')   # 10분 간격으로 수치값 변화
axes[3].plot(df_3f['_time'][::600], df_3f['PM2.5_3f'][::600], color = 'b', label = '3f')
axes[3].legend(loc='best')

axes[4].set_title("PM10.0 24h curve", fontsize = 40)
axes[4].set_xlabel("Time")
axes[4].set_ylabel("PM10.0", fontsize = 30)

axes[4].plot(df_1f['_time'][::600], df_1f['PM10.0_1f'][::600], color = 'r', label ='1f')   # 10분 간격으로 수치값 변화
axes[4].plot(df_3f['_time'][::600], df_3f['PM10.0_3f'][::600], color = 'b', label = '3f')
axes[4].legend(loc='best')
plt.savefig('24h_change.png')
plt.show()

df_3f['CO2_ppm_3f'].mean()


df_mean = pd.DataFrame({'type' : [ 'C02','C0','PM1.0','PM2.5','PM10' ]        ,
                         '1f' : [df_1f['CO2_ppm_1f'].mean() , df_1f['CO_ppm_1f'].mean(), df_1f['PM1.0_1f'].mean(), df_1f['PM2.5_1f'].mean(), df_1f['PM10.0_1f'].mean()]       ,
                        '3f' : [df_3f['CO2_ppm_3f'].mean() , df_3f['CO_ppm_3f'].mean(), df_3f['PM1.0_3f'].mean(), df_3f['PM2.5_3f'].mean(), df_3f['PM10.0_3f'].mean()],
                         })
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches

document = Document()

document.add_paragraph("24 Hours report").bold = True
document.paragraphs[-1].add_run('').italic = True

last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

Text_image_name = '24h_change.png'
document.add_picture(Text_image_name,width= Cm(15), height= Cm(18))


document.add_paragraph("\n평균 표").bold = True
document.paragraphs[-1].add_run('').italic = True

table = document.add_table(rows = 3, cols = 6)
table.style = document.styles['Table Grid']

last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('').bold = True
hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  #가운데 정렬
hdr_cells[1].paragraphs[0].add_run('CO2').bold = True
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[2].paragraphs[0].add_run('C0').bold = True
hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[3].paragraphs[0].add_run('PM 1.0').bold = True
hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[4].paragraphs[0].add_run('PM 2.5').bold = True
hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[5].paragraphs[0].add_run('P, 10,0').bold = True
hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

hdr_cells = table.rows[1].cells
hdr_cells[0].paragraphs[0].add_run('1 층').bold = True
hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[1].paragraphs[0].add_run('{:.2f}'.format(df_1f['CO2_ppm_1f'].mean())).bold = True
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[2].paragraphs[0].add_run('{:.2f}'.format(df_1f['CO2_ppm_1f'].mean())).bold = True
hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[3].paragraphs[0].add_run('{:.2f}'.format(df_1f['PM1.0_1f'].mean())).bold = True
hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[4].paragraphs[0].add_run('{:.2f}'.format(df_1f['PM2.5_1f'].mean())).bold = True
hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[5].paragraphs[0].add_run('{:.2f}'.format(df_1f['PM10.0_1f'].mean())).bold = True
hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

hdr_cells = table.rows[2].cells
hdr_cells[0].paragraphs[0].add_run('2 층').bold = True
hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[1].paragraphs[0].add_run('{:.2f}'.format(df_3f['CO2_ppm_3f'].mean())).bold = True
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[2].paragraphs[0].add_run('{:.2f}'.format(df_3f['CO2_ppm_3f'].mean())).bold = True
hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[3].paragraphs[0].add_run('{:.2f}'.format(df_3f['PM1.0_3f'].mean())).bold = True
hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[4].paragraphs[0].add_run('{:.2f}'.format(df_3f['PM2.5_3f'].mean())).bold = True
hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[5].paragraphs[0].add_run('{:.2f}'.format(df_3f['PM10.0_3f'].mean())).bold = True
hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

document.save('24h_report.docx')
